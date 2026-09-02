"""
Build baseline characteristic tables summarized overall and by outcome.

Continuous variables are summarized as median (Q1, Q3).
Categorical variables are summarized as n (%). 
"""

from typing import Optional, Dict, Any, Iterable
import pandas as pd
from docx import Document

def _infer_var_type(s: pd.Series, cat_threshold: int = 10) -> str:
    s = s.dropna()
    if s.empty:
        return "categorical"

    nunique = s.nunique()

    if nunique == 2:
        return "binary"

    if pd.api.types.is_numeric_dtype(s):
        return "categorical" if nunique <= cat_threshold else "continuous"

    return "categorical"

def _fmt_num(x) -> str:
    if pd.isna(x):
        return ""

    x = float(x)

    if x == 0:
        return "0"

    if round(abs(x), 1) != 0:
        return f"{x:.1f}"
    if round(abs(x), 2) != 0:
        return f"{x:.2f}"

    return "<0.01"

def _median_iqr(x):
    x = pd.to_numeric(x, errors="coerce").dropna()

    if x.empty:
        return ""

    q1, med, q3 = x.quantile([0.25, 0.50, 0.75])

    return f"{_fmt_num(med)} ({_fmt_num(q1)}, {_fmt_num(q3)})"

def _n_pct(n, d):
    if d <= 0:
        return "0 (NA)"

    pct = 100 * n / d

    return f"{n} ({_fmt_num(pct)})"

def _choose_cont_summary(x):
    x = pd.to_numeric(x, errors="coerce").dropna()
    if x.empty:
        return ""
    return _median_iqr(x)

def _get_row_label(var, lvl, level_labels=None):
    """
    Return display label for a row.
    Example:
      level_labels = {"female01": {0: "Male", 1: "Female"}}
    """
    if level_labels and var in level_labels and lvl in level_labels[var]:
        return str(level_labels[var][lvl])
    return f"{var} = {lvl}"


def build_table1(
    df: pd.DataFrame,
    *,
    label_col: int = 0,
    cat_threshold: int = 10,
    level_labels: Optional[Dict[str, Dict[Any, str]]] = None,
    show_all_levels: Optional[Iterable[str]] = None,
) -> pd.DataFrame:
    df = df.copy()

    level_labels = level_labels or {}
    show_all_levels = set(show_all_levels or [])

    # Ensure column names are not blank
    df.columns = [("X" + str(i)) if (c is None or str(c).strip() == "") else str(c) for i, c in enumerate(df.columns)]

    # Identify outcome column
    if label_col < 0 or label_col >= len(df.columns):
        raise ValueError(
            f"label_col={label_col} is out of bounds for "
            f"{len(df.columns)} columns"
        )

    label = df.columns[label_col]
    groups = sorted(df[label].dropna().unique())

    rows = []
    for var in df.columns:
        vtype = _infer_var_type(df[var], cat_threshold)

        if vtype == "continuous":
            row = {"Variable": var, "Level": "", "Type": "Continuous"}
            row["Overall"] = _choose_cont_summary(df[var])
            for g in groups:
                row[f"{label}={g}"] = _choose_cont_summary(df[df[label] == g][var])
            rows.append(row)
            
        elif vtype == "binary":
            levels_to_show = [0, 1] if var in show_all_levels else [1]

            denom_all = int(df[var].notna().sum())

            for lvl in levels_to_show:
                row = {
                    "Variable": var,
                    "Level": _get_row_label(var, lvl, level_labels),
                    "Type": "Binary",
                    "Overall": _n_pct(
                        int((df[var] == lvl).sum()),
                        denom_all,
                    ),
                }

                for g in groups:
                    group_var = df.loc[df[label] == g, var]

                    row[f"{label}={g}"] = _n_pct(
                        int((group_var == lvl).sum()),
                        int(group_var.notna().sum()),
                    )

                rows.append(row)

        else:
            denom_all = int(df[var].notna().sum())
            denom_grp = {g: int(df[df[label] == g][var].notna().sum()) for g in groups}

            for lvl in sorted(df[var].dropna().unique(), key=lambda z: str(z)):
                row = {
                    "Variable": var,
                    "Level": _get_row_label(var, lvl, level_labels),
                    "Type": "Categorical",
                }
                row["Overall"] = _n_pct(int((df[var] == lvl).sum()), denom_all)

                for g in groups:
                    row[f"{label}={g}"] = _n_pct(
                        int((df[df[label] == g][var] == lvl).sum()),
                        denom_grp[g],
                    )

                rows.append(row)

    return pd.DataFrame(rows)

def save_table1_to_word(
    table1: pd.DataFrame,
    out_docx: str,
    *,
    title: str = "Table 1. Descriptive Statistics",
    site_id: Optional[str] = None,
    n_rows: Optional[int] = None,
):
    doc = Document()
    doc.add_heading(title, level=1)

    if site_id:
        doc.add_paragraph(f"Site: {site_id}")
        
    if n_rows is not None:
        doc.add_paragraph(f"N = {n_rows}")

    tbl = doc.add_table(rows=1, cols=table1.shape[1])
    tbl.style = "Table Grid"

    for j, col in enumerate(table1.columns):
        tbl.rows[0].cells[j].text = str(col)

    for _, r in table1.iterrows():
        cells = tbl.add_row().cells
        for j, col in enumerate(table1.columns):
            cells[j].text = "" if pd.isna(r[col]) else str(r[col])

    doc.save(out_docx)
