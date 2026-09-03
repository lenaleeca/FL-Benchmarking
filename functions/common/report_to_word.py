from pathlib import Path

import numpy as np
import pandas as pd

from docx import Document

from functions.common.plot_val_auc import plot_val_auc


# Generic table helpers
def _add_dataframe_table(
    doc,
    df,
    float_cols=None,
    float_fmt="{:.2f}",
):
    """Add a pandas DataFrame to a Word document as a simple grid table."""
    if df is None or df.empty:
        doc.add_paragraph("No data available.")
        return

    float_cols = set(float_cols or [])

    table = doc.add_table(
        rows=1,
        cols=len(df.columns),
    )
    table.style = "Table Grid"

    # Header
    header_cells = table.rows[0].cells

    for j, col in enumerate(df.columns):
        header_cells[j].text = str(col)

    # Data rows
    for _, row in df.iterrows():
        cells = table.add_row().cells

        for j, col in enumerate(df.columns):
            value = row[col]

            if pd.isna(value):
                text = ""

            elif col in float_cols:
                text = float_fmt.format(
                    float(value)
                )

            else:
                text = str(value)

            cells[j].text = text


def _format_ci_string(
    lower,
    upper,
    float_fmt="{:.3f}",
):
    """Format a confidence interval as '(lower, upper)'."""
    if pd.isna(lower) or pd.isna(upper):
        return ""

    return (
        f"({float_fmt.format(float(lower))}, "
        f"{float_fmt.format(float(upper))})"
    )


def _format_test_ci_string(
    test,
    lower,
    upper,
    float_fmt="{:.3f}",
):
    """Format a point estimate and confidence interval."""
    if pd.isna(test):
        return ""

    if pd.isna(lower) or pd.isna(upper):
        return float_fmt.format(
            float(test)
        )

    return (
        f"{float_fmt.format(float(test))} "
        f"({float_fmt.format(float(lower))}, "
        f"{float_fmt.format(float(upper))})"
    )


# Held-out performance helpers
def _reshape_heldout_results_for_word(
    results_df,
):
    """
    Convert wide held-out results into long format:

        site | metric | Test | 95% CI | Test (95% CI)
    """
    columns = [
        "site",
        "metric",
        "Test",
        "95% CI",
        "Test (95% CI)",
    ]

    if results_df is None or results_df.empty:
        return pd.DataFrame(
            columns=columns
        )

    df = results_df.copy()

    id_col = (
        "site_id"
        if "site_id" in df.columns
        else df.columns[0]
    )

    metric_order = [
        "auc",
        "auprc",
        "brier",
        "logloss",
    ]

    rows = []

    for _, row in df.iterrows():
        site_value = row[id_col]

        for metric in metric_order:
            if metric not in df.columns:
                continue

            test_value = row.get(
                metric,
                pd.NA,
            )

            lower_value = row.get(
                f"{metric}_ci_lower",
                pd.NA,
            )

            upper_value = row.get(
                f"{metric}_ci_upper",
                pd.NA,
            )

            rows.append(
                {
                    "site": site_value,
                    "metric": metric,
                    "Test": test_value,
                    "95% CI": _format_ci_string(
                        lower_value,
                        upper_value,
                        float_fmt="{:.3f}",
                    ),
                    "Test (95% CI)": (
                        _format_test_ci_string(
                            test_value,
                            lower_value,
                            upper_value,
                            float_fmt="{:.3f}",
                        )
                    ),
                }
            )

    long_df = pd.DataFrame(
        rows
    )

    if long_df.empty:
        return pd.DataFrame(
            columns=columns
        )

    long_df["site_num"] = pd.to_numeric(
        long_df["site"],
        errors="coerce",
    )

    long_df["metric"] = pd.Categorical(
        long_df["metric"],
        categories=metric_order,
        ordered=True,
    )

    long_df = (
        long_df
        .sort_values(
            by=[
                "site_num",
                "site",
                "metric",
            ],
            na_position="last",
        )
        .drop(
            columns=["site_num"]
        )
    )

    return long_df[
        columns
    ]


def _format_best_round_summary_for_word(
    round_summary_df,
    best_round,
):
    """Return only the selected common best-round row."""
    if (
        round_summary_df is None
        or round_summary_df.empty
        or best_round is None
    ):
        return pd.DataFrame()

    df = round_summary_df.copy()

    if "round" not in df.columns:
        return pd.DataFrame()

    sub = df[
        df["round"].astype(int)
        == int(best_round)
    ].copy()

    if sub.empty:
        return pd.DataFrame()

    keep_cols = [
        col
        for col in [
            "round",
            "mean_val_auc",
            "n_sites",
        ]
        if col in sub.columns
    ]

    return sub[
        keep_cols
    ]

# Cross-site summary helpers
def _format_median_range(
    median,
    minimum,
    maximum,
    float_fmt="{:.3f}",
):
    """Format a median and range as 'median (min, max)'."""
    if pd.isna(median):
        return ""

    return (
        f"{float_fmt.format(float(median))} "
        f"({float_fmt.format(float(minimum))}, "
        f"{float_fmt.format(float(maximum))})"
    )

def _summarise_metrics_across_sites(
    results_df,
    *,
    training_strategy,
    fl_algorithm="",
    source="site held-out",
    metric_specs=None,
    float_fmt="{:.3f}",
):
    """Summarise site-level performance as median (range) across sites."""
    metric_specs = metric_specs or [
        ("auc", "AUC"),
        ("brier", "Brier"),
        ("auprc", "AUPRC"),
    ]

    row = {
        "Training strategy": training_strategy,
        "FL algorithm": fl_algorithm,
        "Source": source,
    }

    for metric, label in metric_specs:
        output_col = f"{label}, median (range)"

        if results_df is None or results_df.empty or metric not in results_df.columns:
            row[output_col] = ""
            continue

        values = pd.to_numeric(
            results_df[metric],
            errors="coerce",
        ).dropna()

        if values.empty:
            row[output_col] = ""
            continue

        row[output_col] = _format_median_range(
            values.median(),
            values.min(),
            values.max(),
            float_fmt=float_fmt,
        )

    if results_df is None or results_df.empty:
        row["n sites"] = ""
    elif "site_id" in results_df.columns:
        row["n sites"] = int(results_df["site_id"].dropna().nunique())
    else:
        row["n sites"] = len(results_df)

    return row


def _format_pooled_metric_row_for_word(
    results_df,
    *,
    training_strategy="Centralised",
    fl_algorithm="",
    source="pooled held-out",
    metric_specs=None,
    float_fmt="{:.3f}",
):
    """
    Format pooled central-model performance for the cross-site summary table.
    """
    metric_specs = metric_specs or [
        ("auc", "AUC"),
        ("brier", "Brier"),
        ("auprc", "AUPRC"),
    ]

    row = {
        "Training strategy": training_strategy,
        "FL algorithm": fl_algorithm,
        "Source": source,
    }

    if results_df is None or results_df.empty:
        for _, label in metric_specs:
            row[
                f"{label}, median (range)"
            ] = ""

        row["n sites"] = ""
        return row

    df = results_df.copy()

    for metric, label in metric_specs:
        output_col = (
            f"{label}, median (range)"
        )

        if metric not in df.columns:
            row[output_col] = ""
            continue

        values = (
            pd.to_numeric(
                df[metric],
                errors="coerce",
            )
            .dropna()
        )

        if values.empty:
            row[output_col] = ""
        else:
            row[output_col] = (
                float_fmt.format(
                    float(values.mean())
                )
            )

    row["n sites"] = "pooled"

    return row


def _format_cross_site_summary_for_word(
    *,
    algorithm_name,
    local_heldout_results_df=None,
    fl_heldout_results_df=None,
    central_pooled_results_df=None,
    include_blank_central_row=True,
    float_fmt="{:.3f}",
):
    
    rows = []

    rows.append(
        _summarise_metrics_across_sites(
            local_heldout_results_df,
            training_strategy="Local",
            fl_algorithm="",
            source="site held-out",
            float_fmt=float_fmt,
        )
    )

    if (
        central_pooled_results_df is not None
        and not central_pooled_results_df.empty
    ):
        rows.append(
            _format_pooled_metric_row_for_word(
                central_pooled_results_df,
                training_strategy="Centralised",
                fl_algorithm="",
                source="pooled held-out",
                float_fmt=float_fmt,
            )
        )

    elif include_blank_central_row:
        rows.append(
            _format_pooled_metric_row_for_word(
                None,
                training_strategy="Centralised",
                fl_algorithm="",
                source="pooled held-out not provided",
                float_fmt=float_fmt,
            )
        )

    rows.append(
        _summarise_metrics_across_sites(
            fl_heldout_results_df,
            training_strategy="FL",
            fl_algorithm=algorithm_name,
            source="site held-out",
            float_fmt=float_fmt,
        )
    )

    df = pd.DataFrame(
        rows
    )

    return df[
        [
            "Training strategy",
            "FL algorithm",
            "Source",
            "AUC, median (range)",
            "Brier, median (range)",
            "AUPRC, median (range)",
            "n sites",
        ]
    ]


# Paired comparison helpers

def _format_difference_ci(
    estimate,
    lower,
    upper,
    decimals=3,
):
    """
    Format a paired difference and confidence interval as:

        estimate (lower, upper)

    Negative zero is suppressed.
    """

    def _format_value(
        value,
    ):
        if pd.isna(value):
            return "NA"

        value = round(
            float(value),
            decimals,
        )

        if value == 0:
            value = 0.0

        return (
            f"{value:.{decimals}f}"
        )

    if pd.isna(estimate):
        return "NA"

    return (
        f"{_format_value(estimate)} "
        f"({_format_value(lower)}, "
        f"{_format_value(upper)})"
    )


def _format_paired_comparison_for_word(
    results_df,
    *,
    include_setting=False,
    decimals=3,
):
    """
    Format paired performance differences for the Word report.

    When include_setting=True, the output includes the analysis setting
    column, such as Within-network or Leave-site-out.
    """
    if (
        results_df is None
        or results_df.empty
    ):
        return pd.DataFrame()

    df = results_df.copy()

    required_cols = {
        "site_id",
        "comparison",
        "auc_difference",
        "auc_ci_lower",
        "auc_ci_upper",
        "brier_difference",
        "brier_ci_lower",
        "brier_ci_upper",
    }

    if include_setting:
        required_cols.add(
            "setting"
        )

    missing_cols = (
        required_cols
        - set(df.columns)
    )

    if missing_cols:
        raise ValueError(
            "Paired comparison dataframe is missing "
            f"required columns: {sorted(missing_cols)}"
        )

    # Display columns

    df["Site"] = (
        df["site_id"]
        .apply(
            lambda x: f"Site {x}"
        )
    )

    df[
        "AUC difference (95% CI)"
    ] = df.apply(
        lambda row: _format_difference_ci(
            row["auc_difference"],
            row["auc_ci_lower"],
            row["auc_ci_upper"],
            decimals=decimals,
        ),
        axis=1,
    )

    df[
        "Brier difference (95% CI)"
    ] = df.apply(
        lambda row: _format_difference_ci(
            row["brier_difference"],
            row["brier_ci_lower"],
            row["brier_ci_upper"],
            decimals=decimals,
        ),
        axis=1,
    )

    # Sort order

    comparison_order = {
        "FL - local": 0,
        "FL - centralised": 1,
        "Centralised - local": 2,
    }

    df["_site_sort"] = pd.to_numeric(
        df["site_id"],
        errors="coerce",
    )

    df["_comparison_order"] = (
        df["comparison"]
        .map(comparison_order)
    )

    sort_cols = [
        "_site_sort",
    ]

    if include_setting:
        setting_order = {
            "Within-network": 0,
            "Leave-site-out": 1,
        }

        df["_setting_order"] = (
            df["setting"]
            .map(setting_order)
        )

        sort_cols.append(
            "_setting_order"
        )

    sort_cols.append(
        "_comparison_order"
    )

    df = df.sort_values(
        sort_cols,
        na_position="last",
    )

    drop_cols = [
        "_site_sort",
        "_comparison_order",
    ]

    if "_setting_order" in df.columns:
        drop_cols.append(
            "_setting_order"
        )

    df = df.drop(
        columns=drop_cols
    )

    # Final table

    if include_setting:
        return df[
            [
                "Site",
                "setting",
                "comparison",
                "AUC difference (95% CI)",
                "Brier difference (95% CI)",
            ]
        ].rename(
            columns={
                "setting": "Setting",
                "comparison": "Comparison",
            }
        )

    return df[
        [
            "Site",
            "comparison",
            "AUC difference (95% CI)",
            "Brier difference (95% CI)",
        ]
    ].rename(
        columns={
            "comparison": "Comparison",
        }
    )


# Stratified report helpers
def _normalise_site_id(
    site_id,
):
    try:
        return int(site_id)
    except Exception:
        return site_id


def _get_sorted_site_ids(
    *dfs,
):
    site_ids = set()

    for df in dfs:
        if (
            df is None
            or df.empty
            or "site_id" not in df.columns
        ):
            continue

        for value in (
            df["site_id"]
            .dropna()
            .unique()
        ):
            site_ids.add(
                _normalise_site_id(
                    value
                )
            )

    return sorted(
        site_ids,
        key=lambda x: (
            int(x)
            if str(x).isdigit()
            else str(x)
        ),
    )


def _format_site_header(
    site_id,
    site_sample_sizes=None,
):
    site_label = (
        f"Site {site_id}"
    )

    if site_sample_sizes is None:
        return site_label

    n = None

    if isinstance(
        site_sample_sizes,
        dict,
    ):
        n = site_sample_sizes.get(
            site_id
        )

        if n is None:
            n = site_sample_sizes.get(
                str(site_id)
            )

        if n is None:
            n = site_sample_sizes.get(
                f"site_{site_id}"
            )

        if n is None:
            n = site_sample_sizes.get(
                f"fl_site_{site_id}"
            )

    if n is None:
        return site_label

    try:
        return (
            f"{site_label}\n"
            f"n={int(n):,}"
        )

    except Exception:
        return (
            f"{site_label}\n"
            f"n={n}"
        )


def _get_auc_ci_value(
    stratified_results_df,
    *,
    site_id,
    model_type,
    dataset,
    float_fmt="{:.3f}",
    show_n=True,
):
    if (
        stratified_results_df is None
        or stratified_results_df.empty
    ):
        return ""

    df = stratified_results_df.copy()

    required_cols = {
        "site_id",
        "model_type",
        "dataset",
        "auc",
    }

    if not required_cols.issubset(
        df.columns
    ):
        return ""

    sub = df[
        (
            df["site_id"]
            .astype(str)
            == str(site_id)
        )
        & (
            df["model_type"]
            .astype(str)
            == str(model_type)
        )
        & (
            df["dataset"]
            .astype(str)
            == str(dataset)
        )
    ]

    if sub.empty:
        return ""

    row = sub.iloc[0]

    auc = row.get(
        "auc",
        pd.NA,
    )

    lower = row.get(
        "auc_ci_lower",
        pd.NA,
    )

    upper = row.get(
        "auc_ci_upper",
        pd.NA,
    )

    n = row.get(
        "n",
        pd.NA,
    )

    n_pos = row.get(
        "n_positive",
        pd.NA,
    )

    n_neg = row.get(
        "n_negative",
        pd.NA,
    )

    one_class = row.get(
        "one_class_outcome",
        False,
    )

    sparse_subgroup = row.get(
        "sparse_subgroup",
        False,
    )

    if pd.isna(auc):
        text = "NA"

    else:
        text = (
            _format_test_ci_string(
                auc,
                lower,
                upper,
                float_fmt=float_fmt,
            )
        )

    if show_n and not pd.isna(n):

        if (
            bool(one_class)
            and not pd.isna(n_pos)
            and not pd.isna(n_neg)
        ):
            text += (
                f"\nn={int(n)}; "
                f"pos={int(n_pos)}; "
                f"neg={int(n_neg)}"
            )

        elif bool(sparse_subgroup):
            text += (
                f"\nn={int(n)}; sparse"
            )

        else:
            text += (
                f"\nn={int(n)}"
            )

    return text


def _get_delta_auc_value(
    sex_auc_difference_df,
    *,
    site_id,
    model_type,
    float_fmt="{:.3f}",
):
    if (
        sex_auc_difference_df is None
        or sex_auc_difference_df.empty
    ):
        return ""

    required_cols = {
        "site_id",
        "model_type",
        "delta_auc_abs",
    }

    if not required_cols.issubset(
        sex_auc_difference_df.columns
    ):
        return ""

    sub = sex_auc_difference_df[
        (
            sex_auc_difference_df[
                "site_id"
            ].astype(str)
            == str(site_id)
        )
        & (
            sex_auc_difference_df[
                "model_type"
            ].astype(str)
            == str(model_type)
        )
    ]

    if sub.empty:
        return ""

    value = sub.iloc[0].get(
        "delta_auc_abs",
        pd.NA,
    )

    if pd.isna(value):
        return ""

    return float_fmt.format(
        float(value)
    )


def _model_display_rows(
    algorithm_name,
):
    """
    Define training-strategy rows used in stratified report tables.
    """
    return [
        {
            "model_type": "local",
            "training_strategy": "Local",
            "fl_algorithm": "",
        },
        {
            "model_type": "central",
            "training_strategy": "Centralised",
            "fl_algorithm": "",
        },
        {
            "model_type": "fl_site",
            "training_strategy": "FL",
            "fl_algorithm": algorithm_name,
        },
    ]


def _set_cell_text(
    cell,
    text,
    bold=False,
):
    cell.text = ""

    paragraph = cell.paragraphs[0]

    run = paragraph.add_run(
        str(text)
    )

    run.bold = bold


def _add_sex_stratified_auc_table(
    doc,
    *,
    stratified_results_df,
    sex_auc_difference_df,
    site_sample_sizes=None,
    algorithm_name,
    float_fmt="{:.3f}",
):
    site_ids = _get_sorted_site_ids(
        stratified_results_df,
        sex_auc_difference_df,
    )

    if not site_ids:
        doc.add_paragraph(
            "No sex-stratified data available."
        )
        return

    doc.add_heading(
        "Table. Discrimination (AUC and 95% CI) "
        "stratified by sex, primary analysis",
        level=2,
    )

    n_cols = (
        3 + len(site_ids)
    )

    table = doc.add_table(
        rows=1,
        cols=n_cols,
    )

    table.style = "Table Grid"

    header_cells = (
        table.rows[0].cells
    )

    headers = [
        "Training\nstrategy",
        "FL\nalgorithm",
        "Sex\nsubgroup",
    ]

    for site_id in site_ids:
        headers.append(
            _format_site_header(
                site_id,
                site_sample_sizes,
            )
        )

    for i, header in enumerate(
        headers
    ):
        _set_cell_text(
            header_cells[i],
            header,
            bold=True,
        )

    for model in _model_display_rows(
        algorithm_name
    ):
        model_type = (
            model["model_type"]
        )

        row_specs = [
            ("Male", "male"),
            ("Female", "female"),
            ("ΔAUC", "delta_auc_abs"),
        ]

        for row_index, (
            label,
            dataset,
        ) in enumerate(
            row_specs
        ):
            cells = (
                table.add_row().cells
            )

            if row_index == 0:
                cells[0].text = (
                    model[
                        "training_strategy"
                    ]
                )
                cells[1].text = (
                    model[
                        "fl_algorithm"
                    ]
                )

            else:
                cells[0].text = ""
                cells[1].text = ""

            cells[2].text = label

            for j, site_id in enumerate(
                site_ids,
                start=3,
            ):
                if (
                    dataset
                    == "delta_auc_abs"
                ):
                    value = (
                        _get_delta_auc_value(
                            sex_auc_difference_df,
                            site_id=site_id,
                            model_type=model_type,
                            float_fmt=float_fmt,
                        )
                    )

                else:
                    value = (
                        _get_auc_ci_value(
                            stratified_results_df,
                            site_id=site_id,
                            model_type=model_type,
                            dataset=dataset,
                            float_fmt=float_fmt,
                        )
                    )

                cells[j].text = value


def _add_age_stratified_auc_table(
    doc,
    *,
    stratified_results_df,
    site_sample_sizes=None,
    algorithm_name,
    float_fmt="{:.3f}",
):
    site_ids = _get_sorted_site_ids(
        stratified_results_df
    )

    if not site_ids:
        doc.add_paragraph(
            "No age-stratified data available."
        )
        return

    doc.add_heading(
        "Table. Discrimination (AUC and 95% CI) "
        "stratified by age group, primary analysis",
        level=2,
    )

    n_cols = (
        3 + len(site_ids)
    )

    table = doc.add_table(
        rows=1,
        cols=n_cols,
    )

    table.style = "Table Grid"

    header_cells = (
        table.rows[0].cells
    )

    headers = [
        "Training\nstrategy",
        "FL\nalgorithm",
        "Age\nsubgroup",
    ]

    for site_id in site_ids:
        headers.append(
            _format_site_header(
                site_id,
                site_sample_sizes,
            )
        )

    for i, header in enumerate(
        headers
    ):
        _set_cell_text(
            header_cells[i],
            header,
            bold=True,
        )

    age_rows = [
        ("18 - 44", "age_18_44"),
        ("45 - 64", "age_45_64"),
        ("65 - 79", "age_65_79"),
        (">= 80", "age_ge_80"),
    ]

    for model in _model_display_rows(
        algorithm_name
    ):
        model_type = (
            model["model_type"]
        )

        for row_index, (
            label,
            dataset,
        ) in enumerate(
            age_rows
        ):
            cells = (
                table.add_row().cells
            )

            if row_index == 0:
                cells[0].text = (
                    model[
                        "training_strategy"
                    ]
                )
                cells[1].text = (
                    model[
                        "fl_algorithm"
                    ]
                )

            else:
                cells[0].text = ""
                cells[1].text = ""

            cells[2].text = label

            for j, site_id in enumerate(
                site_ids,
                start=3,
            ):
                value = (
                    _get_auc_ci_value(
                        stratified_results_df,
                        site_id=site_id,
                        model_type=model_type,
                        dataset=dataset,
                        float_fmt=float_fmt,
                    )
                )

                cells[j].text = value


# Leave-site-out helpers

def _format_leave_site_out_summary_for_word(
    leave_site_out_summary_df,
    *,
    algorithm_name,
    float_fmt="{:.3f}",
):
    """Format the leave-site-out summary table for the Word report."""
    if (
        leave_site_out_summary_df is None
        or leave_site_out_summary_df.empty
    ):
        return pd.DataFrame()

    df = (
        leave_site_out_summary_df
        .copy()
    )

    fl_within_col = (
        f"FL {algorithm_name} within-sample"
    )

    fl_lso_col = (
        f"FL {algorithm_name} leave-site-out"
    )

    rename_map = {
        "excluded_site":
            "Excluded site",

        "local_excluded_site":
            "Local model",

        "central_within_sample":
            "Centralised model within-sample",

        "central_leave_site_out":
            "Centralised model leave-site-out",

        "fl_within_sample":
            fl_within_col,

        "fl_leave_site_out":
            fl_lso_col,

        # Backward-compatible older column names
        "local_xgboost_excluded_site":
            "Local model",

        "centralised_xgboost_within_sample":
            "Centralised model within-sample",

        "centralised_xgboost_leave_site_out":
            "Centralised model leave-site-out",

        "fl_xgboost_histagg_within_sample":
            fl_within_col,

        "fl_xgboost_histagg_leave_site_out":
            fl_lso_col,
    }

    keep_cols = [
        col
        for col in rename_map
        if col in df.columns
    ]

    df = (
        df[keep_cols]
        .rename(
            columns=rename_map
        )
    )

    auc_cols = [
        "Local model",
        "Centralised model within-sample",
        "Centralised model leave-site-out",
        fl_within_col,
        fl_lso_col,
    ]

    for col in auc_cols:
        if col not in df.columns:
            continue

        df[col] = (
            df[col]
            .apply(
                lambda x: (
                    ""
                    if pd.isna(x)
                    else float_fmt.format(
                        float(x)
                    )
                )
            )
        )

    return df


# Secondary-analysis helpers

def _format_secondary_rounds_for_word(
    secondary_best_round_by_site,
):
    if not secondary_best_round_by_site:
        return pd.DataFrame(
            columns=[
                "site_id",
                "site_specific_fl_best_round",
            ]
        )

    rows = []

    for site_id, best_round in (
        secondary_best_round_by_site.items()
    ):
        rows.append(
            {
                "site_id":
                    site_id,

                "site_specific_fl_best_round":
                    int(best_round),
            }
        )

    df = pd.DataFrame(
        rows
    )

    df["site_num"] = (
        pd.to_numeric(
            df["site_id"],
            errors="coerce",
        )
    )

    return (
        df
        .sort_values(
            [
                "site_num",
                "site_id",
            ]
        )
        .drop(
            columns=["site_num"]
        )
    )


# Main report writer

def save_analysis_report_to_word(
    out_docx_path,
    *,
    algorithm_name,
    analysis_label="Analysis Report",
    best_round=None,
    round_summary_df=None,
    metrics_files=None,
    site_sample_sizes=None,

    fl_heldout_results_df=None,
    central_heldout_results_df=None,
    local_heldout_results_df=None,
    central_pooled_results_df=None,

    include_cross_site_summary_table=True,

    secondary_analysis=False,
    secondary_best_round_by_site=None,
    secondary_metrics_files=None,
    secondary_results_df=None,
    secondary_fl_heldout_results_df=None,

    fl_training_report_table_df=None,

    stratified_results_df=None,
    sex_auc_difference_df=None,
    age_auc_summary_df=None,

    leave_site_out_summary_df=None,

    # Paired comparison outputs
    paired_within_site_df=None,
    fidelity_recovery_df=None,

    include_stratified_tables=True,
    include_leave_site_out_table=True,

    val_auc_plot_kwargs=None,
):
    """
    Create the complete Word analysis report.

    The report contains primary held-out evaluation results, cross-site
    summaries, paired within-site performance differences, FL training
    summaries, validation-AUC figures, optional stratified analyses,
    leave-site-out analyses, fidelity/transportability comparisons, and
    optional secondary site-specific-round analyses.
    """

    out_docx_path = Path(
        out_docx_path
    )

    out_docx_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    val_auc_plot_kwargs = dict(
        val_auc_plot_kwargs or {}
    )

    doc = Document()

    doc.add_heading(
        f"FL Analysis - "
        f"{algorithm_name}, "
        f"{analysis_label}",
        level=0,
    )

    # 1. Primary analysis
    doc.add_heading(
        "1. Primary analysis",
        level=1,
    )

    # 1.1 Common best round
    doc.add_heading(
        "1.1 Common best-round selection using the highest mean "
        "validation AUC across participating development sites",
        level=2,
    )

    if best_round is not None:
        paragraph = (
            doc.add_paragraph()
        )

        paragraph.add_run(
            "Selected common best round: "
        ).bold = True

        paragraph.add_run(
            str(int(best_round))
        )

    else:
        doc.add_paragraph(
            "Best round was not provided."
        )

    selected_round_df = (
        _format_best_round_summary_for_word(
            round_summary_df=round_summary_df,
            best_round=best_round,
        )
    )

    if not selected_round_df.empty:
        _add_dataframe_table(
            doc,
            selected_round_df,
            float_cols={
                "mean_val_auc"
            },
            float_fmt="{:.3f}",
        )

    # 1.2 Local performance

    if (
        local_heldout_results_df is not None
        and not local_heldout_results_df.empty
    ):
        doc.add_page_break()

        doc.add_heading(
            "1.2 Performance of each independently trained local model "
            "on its own site held-out evaluation data",
            level=2,
        )

        if (
            "best_iteration"
            in local_heldout_results_df.columns
        ):
            tmp = (
                local_heldout_results_df[
                    [
                        "site_id",
                        "best_iteration",
                    ]
                ]
                .dropna()
                .drop_duplicates()
            )

            if not tmp.empty:
                text = "; ".join(
                    (
                        f"Site {row.site_id}: "
                        f"iteration "
                        f"{int(row.best_iteration)}"
                    )
                    for row in tmp.itertuples(
                        index=False
                    )
                )

                doc.add_paragraph(
                    f"Selected local model iterations: "
                    f"{text}."
                )

        elif (
            "best_epoch"
            in local_heldout_results_df.columns
        ):
            tmp = (
                local_heldout_results_df[
                    [
                        "site_id",
                        "best_epoch",
                    ]
                ]
                .dropna()
                .drop_duplicates()
            )

            if not tmp.empty:
                text = "; ".join(
                    (
                        f"Site {row.site_id}: "
                        f"epoch "
                        f"{int(row.best_epoch)}"
                    )
                    for row in tmp.itertuples(
                        index=False
                    )
                )

                doc.add_paragraph(
                    f"Selected local model epochs: "
                    f"{text}."
                )

        local_word_df = (
            _reshape_heldout_results_for_word(
                local_heldout_results_df
            )
        )

        _add_dataframe_table(
            doc,
            local_word_df,
            float_cols={
                "Test"
            },
            float_fmt="{:.3f}",
        )

    # 1.3 Centralised performance

    if (
        central_heldout_results_df is not None
        and not central_heldout_results_df.empty
    ):
        doc.add_page_break()

        doc.add_heading(
            "1.3 Performance of the central model on each site "
            "held-out evaluation data",
            level=2,
        )

        if (
            "best_iteration"
            in central_heldout_results_df.columns
        ):
            values = (
                central_heldout_results_df[
                    "best_iteration"
                ]
                .dropna()
                .unique()
            )

            if len(values):
                doc.add_paragraph(
                    "Selected central model iteration: "
                    f"{int(values[0])}."
                )

        elif (
            "best_epoch"
            in central_heldout_results_df.columns
        ):
            values = (
                central_heldout_results_df[
                    "best_epoch"
                ]
                .dropna()
                .unique()
            )

            if len(values):
                doc.add_paragraph(
                    "Selected central model epoch: "
                    f"{int(values[0])}."
                )

        central_word_df = (
            _reshape_heldout_results_for_word(
                central_heldout_results_df
            )
        )

        _add_dataframe_table(
            doc,
            central_word_df,
            float_cols={
                "Test"
            },
            float_fmt="{:.3f}",
        )

    # 1.4 FL performance

    if (
        fl_heldout_results_df is not None
        and not fl_heldout_results_df.empty
    ):
        doc.add_page_break()

        doc.add_heading(
            "1.4 Performance of the FL model at the common selected "
            "best round on each site's held-out evaluation data",
            level=2,
        )

        fl_word_df = (
            _reshape_heldout_results_for_word(
                fl_heldout_results_df
            )
        )

        _add_dataframe_table(
            doc,
            fl_word_df,
            float_cols={
                "Test"
            },
            float_fmt="{:.3f}",
        )

    # 1.5 Cross-site summary
    if include_cross_site_summary_table:

        has_cross_site_data = any(
            [
                (
                    local_heldout_results_df
                    is not None
                    and not local_heldout_results_df.empty
                ),
                (
                    fl_heldout_results_df
                    is not None
                    and not fl_heldout_results_df.empty
                ),
                (
                    central_pooled_results_df
                    is not None
                    and not central_pooled_results_df.empty
                ),
            ]
        )

        if has_cross_site_data:
            doc.add_page_break()

            doc.add_heading(
                "1.5 Cross-site held-out performance summary",
                level=2,
            )

            doc.add_paragraph(
                "Values for Local and FL are median (range) across "
                "site-specific held-out performance estimates. "
                "The Centralised row is populated only when pooled "
                "central-model performance is provided by the run script."
            )

            cross_site_summary_df = (
                _format_cross_site_summary_for_word(
                    algorithm_name=algorithm_name,
                    local_heldout_results_df=(
                        local_heldout_results_df
                    ),
                    fl_heldout_results_df=(
                        fl_heldout_results_df
                    ),
                    central_pooled_results_df=(
                        central_pooled_results_df
                    ),
                    include_blank_central_row=True,
                    float_fmt="{:.3f}",
                )
            )

            _add_dataframe_table(
                doc,
                cross_site_summary_df,
                float_cols=None,
                float_fmt="{:.3f}",
            )

    # 1.6 Paired within-site performance differences

    if (
        paired_within_site_df is not None
        and not paired_within_site_df.empty
    ):
        doc.add_page_break()

        doc.add_heading(
            "1.6 Paired within-site performance differences",
            level=2,
        )

        doc.add_paragraph(
            "Differences were calculated using predictions from the same "
            "held-out patients with paired stratified bootstrap 95% "
            "confidence intervals."
        )

        paired_word_df = (
            _format_paired_comparison_for_word(
                paired_within_site_df,
                include_setting=False,
                decimals=3,
            )
        )

        _add_dataframe_table(
            doc,
            paired_word_df,
            float_cols=None,
        )

    # 1.7 FL training time and data transfer

    if (
        fl_training_report_table_df is not None
        and not fl_training_report_table_df.empty
    ):
        doc.add_page_break()

        doc.add_heading(
            "1.7 FL training time and data-transfer summary",
            level=2,
        )

        numeric_float_cols = {
            col
            for col
            in fl_training_report_table_df.columns
            if pd.api.types.is_numeric_dtype(
                fl_training_report_table_df[
                    col
                ]
            )
        }

        _add_dataframe_table(
            doc,
            fl_training_report_table_df,
            float_cols=numeric_float_cols,
            float_fmt="{:.3f}",
        )

    # Figure 1: primary validation AUC

    if metrics_files:
        doc.add_page_break()

        plot_path = (
            out_docx_path.parent
            / (
                f"{algorithm_name}"
                "_primary_val_auc_plot.png"
            )
        )

        plot_val_auc(
            doc=doc,
            metrics_files=metrics_files,
            out_plot_path=plot_path,
            best_round=best_round,
            site_sample_sizes=site_sample_sizes,
            heading=(
                "Figure 1. Site-specific validation AUC across "
                "the communication rounds with a single selected "
                "common best round"
            ),
            **val_auc_plot_kwargs,
        )

    # 1.8 Stratified analysis

    if (
        include_stratified_tables
        and stratified_results_df is not None
        and not stratified_results_df.empty
    ):
        doc.add_page_break()

        doc.add_heading(
            "1.8 Primary stratified analysis",
            level=2,
        )

        _add_sex_stratified_auc_table(
            doc,
            stratified_results_df=(
                stratified_results_df
            ),
            sex_auc_difference_df=(
                sex_auc_difference_df
            ),
            site_sample_sizes=(
                site_sample_sizes
            ),
            algorithm_name=(
                algorithm_name
            ),
            float_fmt="{:.3f}",
        )

        doc.add_page_break()

        _add_age_stratified_auc_table(
            doc,
            stratified_results_df=(
                stratified_results_df
            ),
            site_sample_sizes=(
                site_sample_sizes
            ),
            algorithm_name=(
                algorithm_name
            ),
            float_fmt="{:.3f}",
        )

    # 1.9 Leave-site-out analysis

    if (
        include_leave_site_out_table
        and leave_site_out_summary_df is not None
        and not leave_site_out_summary_df.empty
    ):
        doc.add_page_break()

        doc.add_heading(
            "1.9 Primary leave-site-out analysis",
            level=2,
        )

        leave_site_out_word_df = (
            _format_leave_site_out_summary_for_word(
                leave_site_out_summary_df,
                algorithm_name=algorithm_name,
                float_fmt="{:.3f}",
            )
        )

        _add_dataframe_table(
            doc,
            leave_site_out_word_df,
            float_cols=None,
            float_fmt="{:.3f}",
        )

    # 1.10 Fidelity / transportability

    if (
        fidelity_recovery_df is not None
        and not fidelity_recovery_df.empty
    ):
        doc.add_page_break()

        doc.add_heading(
            "1.10 Fidelity and transportability",
            level=2,
        )

        doc.add_paragraph(
            "Shared models were compared with the site's independently "
            "trained local model under within-network and leave-site-out "
            "settings."
        )

        fidelity_word_df = (
            _format_paired_comparison_for_word(
                fidelity_recovery_df,
                include_setting=True,
                decimals=3,
            )
        )

        _add_dataframe_table(
            doc,
            fidelity_word_df,
            float_cols=None,
        )

    # 2. Secondary analysis

    if secondary_analysis:
        doc.add_page_break()

        doc.add_heading(
            "2. Secondary analysis",
            level=1,
        )

        # 2.1 Site-specific best rounds

        secondary_rounds_df = (
            _format_secondary_rounds_for_word(
                secondary_best_round_by_site
            )
        )

        if not secondary_rounds_df.empty:
            doc.add_heading(
                "2.1 Site-specific selected FL best round by "
                "site validation AUC",
                level=2,
            )

            _add_dataframe_table(
                doc,
                secondary_rounds_df,
            )

        # 2.2 Secondary held-out performance

        if (
            secondary_fl_heldout_results_df
            is not None
            and not secondary_fl_heldout_results_df.empty
        ):
            doc.add_page_break()

            doc.add_heading(
                "2.2 Performance of each site's FL best round model "
                "on that site's held-out evaluation data",
                level=2,
            )

            secondary_fl_word_df = (
                _reshape_heldout_results_for_word(
                    secondary_fl_heldout_results_df
                )
            )

            _add_dataframe_table(
                doc,
                secondary_fl_word_df,
                float_cols={
                    "Test"
                },
                float_fmt="{:.3f}",
            )

            # 2.3 Cross-site secondary summary

            if include_cross_site_summary_table:
                doc.add_page_break()

                doc.add_heading(
                    "2.3 Cross-site held-out performance summary "
                    "using site-specific FL best rounds",
                    level=2,
                )

                doc.add_paragraph(
                    "Values are median (range) across site-specific held-out "
                    "performance estimates. The FL row uses each site's "
                    "own selected best round."
                )

                secondary_summary_df = (
                    _format_cross_site_summary_for_word(
                        algorithm_name=algorithm_name,
                        local_heldout_results_df=(
                            local_heldout_results_df
                        ),
                        fl_heldout_results_df=(
                            secondary_fl_heldout_results_df
                        ),
                        central_pooled_results_df=(
                            central_pooled_results_df
                        ),
                        include_blank_central_row=True,
                        float_fmt="{:.3f}",
                    )
                )

                _add_dataframe_table(
                    doc,
                    secondary_summary_df,
                    float_cols=None,
                    float_fmt="{:.3f}",
                )

        # Figure 2: secondary validation AUC

        if secondary_metrics_files:
            doc.add_page_break()

            secondary_plot_path = (
                out_docx_path.parent
                / (
                    f"{algorithm_name}"
                    "_secondary_site_best_val_auc_plot.png"
                )
            )

            plot_val_auc(
                doc=doc,
                metrics_files=secondary_metrics_files,
                out_plot_path=secondary_plot_path,
                best_round_by_site=(
                    secondary_best_round_by_site
                ),
                site_sample_sizes=(
                    site_sample_sizes
                ),
                heading=(
                    "Figure 2. Site-specific validation AUC across "
                    "the communication rounds with site-specific "
                    "selected best rounds"
                ),
                **val_auc_plot_kwargs,
            )

    doc.save(
        out_docx_path
    )