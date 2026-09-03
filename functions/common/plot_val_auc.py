"""
Create validation AUC plots across federated learning communication rounds.

The function reads site-level metrics CSV files, plots validation AUC trajectories,
and optionally annotates a common or site-specific selected FL round. The figure
is saved in the requested formats and can also be inserted directly into a Word
report.
"""

from pathlib import Path
import math

import matplotlib.pyplot as plt
import pandas as pd
from docx.shared import Inches

DEFAULT_FIGSIZE = (7, 2.8)
DEFAULT_DPI = 300

DEFAULT_X_LABEL = "FL communication round"
DEFAULT_Y_LABEL = "Site-specific validation AUC"

DEFAULT_PALETTE = [
    "#0072B2",
    "#D55E00",
    "#009E73",
    "#E69F00",
    "#CC79A7",
    "#56B4E9",
    "#F0E442",
    "#000000",
]


def _lookup_site_value(mapping, site_id):
    if not mapping:
        return None

    candidates = [site_id, str(site_id)]

    if str(site_id).isdigit():
        candidates.append(int(str(site_id)))

    for key in candidates:
        if key in mapping:
            return mapping[key]

    return None


def _read_validation_auc(path):
    """
    Read validation AUC values from one metrics CSV.

    Expected columns:
        round, dataset, metric, value
    """
    path = Path(path)

    if not path.exists():
        raise FileNotFoundError(
            f"Metrics file not found: {path}"
        )

    df = pd.read_csv(path)

    required = {"round", "dataset", "metric", "value"}
    missing = required - set(df.columns)

    if missing:
        raise ValueError(
            f"{path} is missing required columns: {sorted(missing)}"
        )
        
    sub = df[
        df["dataset"].astype(str).str.lower().eq("val")
        & df["metric"].astype(str).str.lower().eq("auc")
    ][["round", "value"]].copy()

    sub["round"] = pd.to_numeric(sub["round"], errors="coerce")
    sub["value"] = pd.to_numeric(sub["value"], errors="coerce")
    sub = sub.dropna().sort_values("round")

    if sub.empty:
        raise ValueError(f"No validation AUC rows found in: {path}")

    sub["round"] = sub["round"].astype(int)

    return sub


def _auto_x_tick_interval(min_round, max_round):
    span = max_round - min_round

    if span <= 20:
        return 1
    if span <= 50:
        return 5
    if span <= 100:
        return 10
    if span <= 250:
        return 25
    if span <= 500:
        return 50

    return 100


def _auto_legend_ncol(
    n_items,
    max_legend_rows=2,
):
    if n_items <= 0:
        return 1

    return max(1, math.ceil(n_items / max(1, max_legend_rows)))


def plot_val_auc(
    doc,
    metrics_files,
    out_plot_path,
    *,
    best_round=None,
    best_round_by_site=None,
    site_sample_sizes=None,

    # Word
    heading="Validation AUC across rounds by site",
    width_inches=6.5,

    # Figure
    figsize=DEFAULT_FIGSIZE,
    dpi=DEFAULT_DPI,
    formats=("png", "pdf"),

    # Axes
    x_label=DEFAULT_X_LABEL,
    y_label=DEFAULT_Y_LABEL,
    x_tick_interval=None,
    y_axis_min=None,
    y_axis_max=None,
    y_pad_ratio=0.08,

    # Plot appearance
    show_grid=True,
    line_width=2.2,
    marker_size=80,
    palette=DEFAULT_PALETTE,

    # Fonts
    axis_label_fontsize=7,
    tick_label_fontsize=8,
    legend_fontsize=6.8,

    # Legend
    legend_ncol=None,
    max_legend_rows=2,
    legend_bottom_y=0.015,
    legend_columnspacing=0.9,
    legend_handlelength=2.5,
    legend_labelspacing=0.35,

    # Margins
    bottom_margin=0.28,
    left_margin=0.10,
    right_margin=0.98,
    top_margin=0.95,
):

    out_plot_path = Path(out_plot_path)
    out_plot_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    if best_round is not None:
        best_round = int(best_round)

    fig, ax = plt.subplots(figsize=figsize)

    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    all_values = []
    all_rounds = []
    plotted_any = False

    for i, (site_id, metrics_path) in enumerate((metrics_files or {}).items()):
        metrics_path = Path(metrics_path)

        if not metrics_path.exists():
            print(f"Skipping missing metrics file: {metrics_path}")
            continue

        sub = _read_validation_auc(metrics_path)
        color = palette[i % len(palette)]
        n_site = _lookup_site_value(site_sample_sizes, site_id)

        if n_site is not None:
            legend_label = (
                f"Site {site_id} "
                f"(N = {int(n_site):,})"
            )
        else:
            legend_label = f"Site {site_id}"

        ax.plot(
            sub["round"],
            sub["value"],
            linewidth=line_width,
            color=color,
            label=legend_label,
            zorder=3,
        )


        selected_round = _lookup_site_value(
            best_round_by_site,
            site_id,
        )

        if selected_round is not None:
            selected_round = int(selected_round)
            selected = sub[sub["round"] == selected_round]

            if selected.empty:
                print(
                    f"Warning: selected round {selected_round} "
                    f"for Site {site_id} was not found."
                )
            else:
                selected_auc = float(
                    selected.iloc[0]["value"]
                )

                ax.scatter(
                    [selected_round],
                    [selected_auc],
                    s=marker_size,
                    marker="D",
                    facecolor=color,
                    edgecolor="black",
                    linewidth=1.4,
                    zorder=9,
                )

        all_values.extend(sub["value"].astype(float))
        all_rounds.extend(sub["round"].astype(int))

        plotted_any = True

    doc.add_heading(heading, level=2)

    if not plotted_any:
        doc.add_paragraph(
            "No validation AUC plot could be created."
        )
        plt.close(fig)
        return []


    if best_round is not None:
        ax.axvline(
            x=best_round,
            color="red",
            linestyle="--",
            linewidth=2.2,
            alpha=0.9,
            label=(
                "Common selected best round "
                f"= {best_round}"
            ),
            zorder=6,
        )
        all_rounds.append(best_round)

    # X axis
    min_round = min(all_rounds)
    max_round = max(all_rounds)
    if not x_tick_interval or x_tick_interval <= 0:
        x_tick_interval = _auto_x_tick_interval(
            min_round,
            max_round,
        )

    x_tick_interval = int(x_tick_interval)

    start_tick = (
        min_round // x_tick_interval
    ) * x_tick_interval

    xticks = list(
        range(
            start_tick,
            max_round + 1,
            x_tick_interval,
        )
    )

    xticks.extend([min_round, max_round])

    x_pad = (
        0.5
        if min_round == max_round
        else max(
            0.5,
            (max_round - min_round) * 0.02,
        )
    )

    ax.set_xticks(sorted(set(xticks)))
    ax.set_xlim(
        min_round - x_pad,
        max_round + x_pad,
    )

    # Y axis
    observed_min = min(all_values)
    observed_max = max(all_values)
    y_pad = (
        0.02
        if observed_min == observed_max
        else (observed_max - observed_min) * y_pad_ratio
    )
    
    y_lower = (
        max(0, observed_min - y_pad)
        if y_axis_min is None
        else float(y_axis_min)
    )

    y_upper = (
        min(1, observed_max + y_pad)
        if y_axis_max is None
        else float(y_axis_max)
    )

    ax.set_ylim(y_lower, y_upper)
    ax.set_xlabel(x_label, fontsize=axis_label_fontsize)

    ax.set_ylabel(
        y_label,
        fontsize=axis_label_fontsize,
        labelpad=10,
    )

    ax.tick_params(
        axis="both",
        labelsize=tick_label_fontsize,
    )

    if show_grid:
        ax.grid(
            True,
            axis="y",
            alpha=0.2,
        )

    handles, labels = ax.get_legend_handles_labels()

    if legend_ncol is None:
        legend_ncol = _auto_legend_ncol(
            len(labels),
            max_legend_rows,
        )

    legend = fig.legend(
        handles,
        labels,
        loc="lower center",
        bbox_to_anchor=(0.55, legend_bottom_y),
        ncol=legend_ncol,
        frameon=False,
        fontsize=legend_fontsize,
        columnspacing=legend_columnspacing,
        handlelength=legend_handlelength,
        labelspacing=legend_labelspacing,
        borderaxespad=0.0,
    )

    legend.set_zorder(1000)

    fig.subplots_adjust(
        left=left_margin,
        right=right_margin,
        bottom=bottom_margin,
        top=top_margin,
    )

    saved_paths = []
    base_path = out_plot_path.with_suffix("")

    for fmt in formats:
        fmt = str(fmt).lower().lstrip(".")
        save_path = base_path.with_suffix(f".{fmt}")
        save_kwargs = {"facecolor": "white"}

        if fmt in {"png", "jpg", "jpeg", "tif", "tiff"}:
            save_kwargs["dpi"] = dpi

        fig.savefig(
            save_path,
            **save_kwargs,
        )
        saved_paths.append(save_path)

    plt.close(fig)

    png_path = next(
        (
            path
            for path in saved_paths
            if path.suffix.lower() == ".png"
        ),
        None,
    )

    if png_path is None:
        raise ValueError(
            "PNG must be included in formats because "
            "the Word report inserts the PNG version."
        )

    doc.add_picture(
        str(png_path),
        width=Inches(width_inches),
    )

    return saved_paths